import io
import json
import os
import tempfile
import time

import docx
import openai
import requests
from docling.document_converter import DocumentConverter
from docx.shared import Cm, Pt
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from fpdf import FPDF
from loguru import logger
from pydantic import BaseModel, Field

load_dotenv()

WIKIDATA_SPARQL_URL = os.getenv("WIKIDATA_SPARQL_URL", "https://query.wikidata.org/sparql")
USER_AGENT = os.getenv("USER_AGENT", "AcademicCVGenerator")
LLM_PROVIDER = os.getenv("LLM_PROVIDER", "campusai")

if LLM_PROVIDER == "campusai":
    LLM_BASE_URL = os.getenv(
        "CAMPUSAI_BASE_URL", "https://api.campusai.compute.dtu.dk/v1"
    )
    LLM_MODEL = os.getenv("CAMPUSAI_MODEL", "Gemma 4")
    LLM_API_KEY = os.getenv("CAMPUSAI_API_KEY", "your-campusai-key")
    logger.info("Using CampusAI as LLM provider")
else:
    LLM_BASE_URL = os.getenv("LMSTUDIO_BASE_URL", "http://localhost:1234/v1")
    LLM_MODEL = os.getenv("LMSTUDIO_MODEL", "google/gemma-4-26b-a4b")
    LLM_API_KEY = os.getenv("LMSTUDIO_API_KEY", "lm-studio")
    logger.info("Using LM Studio as LLM provider")


app = FastAPI(
    title="Academic CV Generator API",
    version="0.1.0",
    description="Builds a starter academic CV from Wikidata and simple user preferences.",
)


class ChatMessage(BaseModel):
    role: str = Field(..., examples=["system", "user"])
    content: str = Field(..., examples=["You are an academic CV assistant."])


def get_llm_client() -> openai.OpenAI:
    """Initializes and returns an OpenAI client configured for the selected LLM provider."""
    return openai.OpenAI(
        api_key=LLM_API_KEY,
        base_url=LLM_BASE_URL,
    )


def call_llm(messages: list[ChatMessage]) -> str:
    """Calls the configured LLM provider with the given messages and returns the response content."""
    client = get_llm_client()

    try:
        completion = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[message.model_dump() for message in messages],
        )
    except openai.AuthenticationError as exc:
        raise HTTPException(
            status_code=502,
            detail="LLM authentication failed. Check provider and API key configuration.",
        ) from exc

    if not completion.choices:
        raise HTTPException(
            status_code=502, detail="LLM response did not include any choices"
        )

    message = completion.choices[0].message
    if message is None or message.content is None:
        raise HTTPException(
            status_code=502, detail="LLM response did not include message content"
        )

    return message.content


async def extract_text_from_pdf(file: UploadFile) -> str:
    """Extracts text from an uploaded PDF file using Docling."""
    try:
        if not file.filename.endswith(".pdf"):
            raise HTTPException(status_code=400, detail="Only PDF files are supported")

        logger.info(f"Processing uploaded file: {file.filename}")
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            content = await file.read()
            temp_pdf.write(content)
            temp_pdf_path = temp_pdf.name

        converter = DocumentConverter()
        doc = converter.convert(temp_pdf_path).document

        text = doc.export_to_markdown()
        logger.info("PDF converted to markdown successfully with Docling")

        os.unlink(temp_pdf_path)

        return text

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


def execute_sparql(query, max_retries=3):
    """Executes a SPARQL query against Wikidata with automatic retries."""
    headers = {
        "Accept": "application/json",
        "User-Agent": USER_AGENT,
    }

    for attempt in range(max_retries):
        try:
            response = requests.get(
                WIKIDATA_SPARQL_URL, params={"query": query}, headers=headers
            )

            if response.status_code in [429, 500, 502, 503, 504]:
                wait_time = 5 * (attempt + 1)
                logger.warning(
                    f"Wikidata error {response.status_code}. Retrying in {wait_time}s..."
                )
                time.sleep(wait_time)
                continue

            response.raise_for_status()

            return response.json().get("results", {}).get("bindings", [])

        except requests.exceptions.RequestException as e:
            logger.warning(f"Network error: {e}. Retrying...")
            time.sleep(5)

    logger.error("Max retries reached. Returning empty list to prevent crash.")
    return []


def format_for_llm(raw_sparql: dict) -> dict:
    """Cleans and formats raw SPARQL results into a structured dictionary for that is more suitable for LLM input."""
    cleaned_data = {"name": raw_sparql.get("name", "Unknown")}

    for section, items in raw_sparql.items():
        if section == "name":
            continue

        cleaned_list = []
        for item in items:
            clean_item = {}
            for key, data_dict in item.items():
                value = data_dict.get("value", "")

                if value.startswith("http://www.wikidata.org/entity/"):
                    continue

                if "T00:00:00Z" in value:
                    value = value.replace("T00:00:00Z", "")

                clean_item[key] = value

            if clean_item and clean_item not in cleaned_list:
                cleaned_list.append(clean_item)

        cleaned_data[section] = cleaned_list

    return cleaned_data


def get_researcher_data(qid: str):
    """Fetches and formats researcher data from Wikidata for a given QID."""
    queries = {
        "education": f"""
            SELECT ?name ?educationLabel ?degreeLabel ?eduStart ?eduEnd WHERE {{
            BIND(wd:{qid} AS ?researcher)
            ?researcher rdfs:label ?name . FILTER(LANG(?name) = "en")
            ?researcher p:P69 ?eduStatement . ?eduStatement ps:P69 ?education .
            OPTIONAL {{ ?eduStatement pq:P512 ?degree . }}
            OPTIONAL {{ ?eduStatement pq:P580 ?eduStart . }}
            OPTIONAL {{ ?eduStatement pq:P582 ?eduEnd . }}
            SERVICE wikibase:label {{ bd:serviceParam wikibase:language "en". }}
            }} ORDER BY DESC(?eduEnd) DESC(?eduStart)
        """,
        "employment": f"""
            SELECT ?name ?employerLabel ?roleLabel ?empStart ?empEnd WHERE {{
            BIND(wd:{qid} AS ?researcher)
            ?researcher rdfs:label ?name . FILTER(LANG(?name) = "en")
            ?researcher p:P108 ?empStatement . ?empStatement ps:P108 ?employer .
            OPTIONAL {{ ?empStatement pq:P39 ?role . }}
            OPTIONAL {{ ?empStatement pq:P580 ?empStart . }}
            OPTIONAL {{ ?empStatement pq:P582 ?empEnd . }}
            SERVICE wikibase:label {{ bd:serviceParam wikibase:language "en". }}
            }} ORDER BY DESC(?empStart)
        """,
        "awards": f"""
            SELECT ?name ?awardLabel ?awardDate WHERE {{
            BIND(wd:{qid} AS ?researcher)
            ?researcher rdfs:label ?name . FILTER(LANG(?name) = "en")
            ?researcher p:P166 ?awardStatement . ?awardStatement ps:P166 ?award .
            OPTIONAL {{ ?awardStatement pq:P585 ?awardDate . }}
            SERVICE wikibase:label {{ bd:serviceParam wikibase:language "en". }}
            }} ORDER BY DESC(?awardDate)
        """,
        "supervision": f"""
            SELECT ?role ?personLabel WHERE {{
            BIND(wd:{qid} AS ?researcher)
            {{ ?researcher wdt:P185 ?person . BIND("Supervised Student" AS ?role) }}
            UNION
            {{ ?researcher wdt:P184 ?person . BIND("Academic Advisor" AS ?role) }}
            SERVICE wikibase:label {{ bd:serviceParam wikibase:language "en". }}
            }}
        """,
        "collaborations": f"""
            SELECT DISTINCT ?coauthorLabel WHERE {{
            BIND(wd:{qid} AS ?researcher)
            ?work wdt:P50 ?researcher .
            ?work wdt:P50 ?coauthor .
            FILTER(?coauthor != ?researcher)
            SERVICE wikibase:label {{ bd:serviceParam wikibase:language "en". }}
            }} LIMIT 50
        """,
        "track_record": f"""
            SELECT DISTINCT ?workLabel ?date WHERE {{
            BIND(wd:{qid} AS ?researcher)
            ?work wdt:P50 ?researcher .
            OPTIONAL {{ ?work wdt:P577 ?date . }}
            SERVICE wikibase:label {{ bd:serviceParam wikibase:language "en". }}
            }} ORDER BY DESC(?date) LIMIT 50
        """,
    }

    logger.info(f"Querying wikidata for QID: {qid}")
    results = {}
    for key, q in queries.items():
        results[key] = execute_sparql(q)
        logger.info(f"Fetched {len(results[key])} results for {key}")
        time.sleep(3)  # avoid hitting rate limits on wikidata

    name = "Unknown"
    for res in results.values():
        if res and "name" in res[0]:
            name = res[0]["name"]["value"]
            break

    raw_data = {
        "name": name,
        "education": results["education"],
        "employment": results["employment"],
        "awards": results["awards"],
        "supervision": results["supervision"],
        "collaborations": results["collaborations"],
        "track_record": results["track_record"],
    }

    return format_for_llm(raw_data)


def get_template_instructions(filepath: str) -> str:
    """Reads the raw text from the Word template to feed to the LLM."""
    try:
        doc = docx.Document(filepath)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception:
        return ""


@app.post("/api/v1/research/{wikidata_qid}")
def research(wikidata_qid: str):
    '''Fetches and formats researcher data from Wikidata for a given QID.'''
    profile = get_researcher_data(wikidata_qid)
    return profile


@app.get("/health")
def health() -> dict[str, str]:
    """Simple health check endpoint."""
    return {"status": "ok"}


@app.post("/api/v1/generate/{wikidata_qid}")
async def generate_cv(
    wikidata_qid: str, format: str = "docx", previous_cv: UploadFile = File(None)
):
    '''Generates a CV based on Wikidata profile and optional previous CV.'''
    if wikidata_qid == "Q20980928" and os.path.exists(
        "research_Q20980928_results.json"
    ):
        profile = json.load(open("research_Q20980928_results.json", "r"))
    else:
        profile = get_researcher_data(wikidata_qid)

    if not profile:
        raise HTTPException(
            status_code=404, detail="Researcher not found or no data available."
        )
    logger.info(
        f"Generating CV for {profile['name']} (QID: {wikidata_qid}) with previous CV: {previous_cv.filename if previous_cv else 'None'}"
    )
    # logger.debug(f"Wikidata profile data: {profile}")

    previous_cv_text = ""
    if previous_cv:
        try:
            previous_cv_text = await extract_text_from_pdf(previous_cv)
            logger.info(
                f"Extracted text from previous CV: {len(previous_cv_text)} characters"
            )
        except Exception as e:
            logger.error(f"Error extracting text from previous CV: {e}")
            raise HTTPException(
                status_code=400,
                detail=f"Failed to extract text from uploaded CV: {str(e)}",
            )

    template_instructions = get_template_instructions("dff-cv-template.docx")

    format_instructions = {
        "markdown": "Format the output using standard Markdown (## headers, **bold**, *italics*, bullet points).",
        "latex": "Format the output as a COMPLETE, compilable LaTeX document using the \\documentclass{article} class. Do not use markdown.",
        "docx": "Format the output clearly as plain text for a docx file. Headers can be marked with #.",
        "pdf": "Format the output clearly as plain text for a pdf file. Headers can be marked with #.",
    }

    selected_instruction = format_instructions.get(
        format.lower(), format_instructions[f"{format}"]
    )

    system_instruction = f"""You are an expert consultant for the DFF 2026 call. 
    Draft a CV that includes these EXACT headers:
    1. CV (as a title)
    2. Family name, First name(s)
    3. Current position(s)
    4. Previous positions
    5. Education
    6. Career breaks
    7. Research statement
    8. Personal context
    9. Grants and awards
    10. Supervision, teaching and research leadership
    11. Collaborations and teamwork
    12. Contributions to the research community
    13. Contributions to the wider society

    Constraints: No links, no h-index, 12pt font style content, the final document should be max 3 pages. {selected_instruction}"""

    user_prompt = f"""
    Draft a DFF 2026 narrative CV using this Wikidata profile:
    
    Name: {profile.get("name", "")}
    Education: {json.dumps(profile.get("education", []), indent=2)}
    Employment: {json.dumps(profile.get("employment", []), indent=2)}
    Awards: {json.dumps(profile.get("awards", []), indent=2)}
    Supervision: {json.dumps(profile.get("supervision", []), indent=2)}
    Collaborations: {json.dumps(profile.get("collaborations", []), indent=2)}
    Track Record: {json.dumps(profile.get("track_record", []), indent=2)}

    Here is the official DFF CV Template and its instructions. Read the bracketed text carefully (e.g., length limits) and write the content for each section accordingly:
    {template_instructions}
    """

    if previous_cv_text:
        user_prompt += f"\n\nAdditionally, incorporate relevant narrative details from the applicant's previous CV:\n{previous_cv_text}"

    messages = [
        ChatMessage(role="system", content=system_instruction),
        ChatMessage(role="user", content=user_prompt),
    ]

    llm_text = call_llm(messages)

    file_stream = io.BytesIO()
    filename = f"DFF_CV_{profile.get('name', 'Draft').replace(' ', '_')}"

    if format == "markdown":
        file_stream.write(llm_text.encode("utf-8"))
        media_type = "text/markdown"
        filename += ".md"

    elif format == "latex":
        file_stream.write(llm_text.encode("utf-8"))
        media_type = "application/x-tex"
        filename += ".tex"

    elif format == "docx":
        doc = docx.Document()

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5

        for line in llm_text.split("\n"):
            line = line.strip()
            if not line:
                continue

            # Check if the line starts with a '#' to signify a header
            if line.startswith("#"):
                clean_header = line.lstrip("#").strip()
                doc.add_heading(clean_header, level=3)
            else:
                doc.add_paragraph(line)

        doc.save(file_stream)
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        filename += ".docx"

    elif format == "pdf":
        pdf = FPDF()
        pdf.set_margins(left=20, top=20, right=20)
        pdf.add_page()

        for line in llm_text.split("\n"):
            line = line.strip()
            if not line:
                pdf.ln(5)
                continue

            # Check if the line starts with a '#' to signify a header
            if line.startswith("#"):
                clean_header = line.lstrip("#").strip()
                safe_line = clean_header.encode("latin-1", "replace").decode("latin-1")

                pdf.set_font("Times", "B", size=14)
                pdf.multi_cell(0, 10, txt=safe_line)

                pdf.set_font("Times", size=12)
            else:
                safe_line = line.encode("latin-1", "replace").decode("latin-1")
                pdf.multi_cell(0, 7, txt=safe_line)

        pdf_bytes = pdf.output(dest="S")
        if isinstance(pdf_bytes, str):
            pdf_bytes = pdf_bytes.encode("latin-1")

        file_stream.write(pdf_bytes)
        media_type = "application/pdf"
        filename += ".pdf"

    else:
        raise HTTPException(status_code=400, detail="Invalid format requested.")

    file_stream.seek(0)
    return StreamingResponse(
        file_stream,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
